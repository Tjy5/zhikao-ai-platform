#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
基于人类逻辑的题目提取器服务
实现简单、高效、准确的题目提取逻辑
"""

from docx import Document
from docx.oxml.ns import nsmap
import re
import os
import uuid
from typing import List, Dict, Any, Optional, Tuple
from PIL import Image
import io
import logging

logger = logging.getLogger(__name__)


class HumanLogicQuestionExtractor:
    """基于人类逻辑的题目提取器"""
    
    def __init__(self):
        self.section_pattern = r'^[一二三四五六]、'
        self.question_pattern = r'^\d+$'
        
    def extract_questions(self, docx_path: str) -> Dict[str, Any]:
        """主提取函数"""
        try:
            doc = Document(docx_path)
            logger.info(f"开始提取题目，文档路径: {docx_path}")
            
            # 第1步：找到所有题型分界点
            sections = self.find_sections(doc)
            logger.info(f"找到 {len(sections)} 个题型")
            
            # 第2步：在每个题型内找题目
            total_questions = 0
            total_images = 0
            
            for section in sections:
                section['questions'] = self.find_questions_in_section(doc, section)
                logger.info(f"{section['name']}: {len(section['questions'])}道题")
                total_questions += len(section['questions'])
            
            # 第3步：提取每道题的完整内容（按数据集分组合并材料）
            all_questions = []
            for section in sections:
                section_questions = section['questions']

                # 计算每个组（（一）（二）（三）…）的材料前言，仅在组内共享
                preface_by_group = {}
                first_num_idx_by_group = {}
                if section_questions:
                    from collections import defaultdict
                    grouped = defaultdict(list)
                    for q in section_questions:
                        gid = q.get('group_id', 0)
                        grouped[gid].append(q)
                    for gid, qs in grouped.items():
                        first_q = min(qs, key=lambda x: x.get('number_index', x['start']))
                        number_index = first_q.get('number_index', first_q['start'])
                        if first_q['start'] < number_index:
                            preface_by_group[gid] = self.extract_question_content(doc, {
                                'start': first_q['start'],
                                'end': number_index
                            })
                        first_num_idx_by_group[gid] = number_index

                for question in section_questions:
                    raw_content = self.extract_question_content(doc, question)
                    content = raw_content
                    gid = question.get('group_id', 0)
                    first_idx = first_num_idx_by_group.get(gid)
                    is_first_in_group = (first_idx is not None and question.get('number_index') == first_idx)
                    # 非本组首题需要叠加本组材料
                    if gid in preface_by_group and not is_first_in_group:
                        pf = preface_by_group[gid]
                        content = {
                            'paragraphs': list(pf['paragraphs']) + raw_content['paragraphs'],
                            'paragraph_count': pf['paragraph_count'] + raw_content['paragraph_count'],
                            'total_images': pf['total_images'] + raw_content['total_images'],
                            'total_text_length': pf['total_text_length'] + raw_content['total_text_length'],
                        }
                    question_data = {
                        'number': question['number'],
                        'section': section['name'],
                        'number_index': question.get('number_index'),
                        'group_id': question.get('group_id'),
                        'content': content,
                        'paragraph_range': f"{question['start']}-{question['end']}"
                    }
                    all_questions.append(question_data)
                    # 注意：此处累计的是原始题目图片数，用于整体校验；不重复累计组前言
                    total_images += raw_content.get('total_images', 0)
            
            # 验证结果
            validation = self.validate_extraction_results(all_questions, total_questions, total_images)
            
            return {
                'success': True,
                'total_questions': total_questions,
                'total_images': total_images,
                'sections': {s['name'][:2]: {
                    'name': s['name'],
                    'count': len(s['questions']),
                    'questions': [q['number'] for q in s['questions']],
                    'total_images': sum(self.extract_question_content(doc, q).get('total_images', 0) for q in s['questions'])
                } for s in sections},
                'questions': all_questions,
                'validation': validation
            }
            
        except Exception as e:
            logger.error(f"题目提取失败: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'total_questions': 0,
                'total_images': 0,
                'sections': {},
                'questions': [],
                'validation': {'issues': [f'提取过程错误: {str(e)}']}
            }
    
    def find_sections(self, doc: Document) -> List[Dict[str, Any]]:
        """找题型边界"""
        sections = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if re.match(self.section_pattern, text):
                sections.append({
                    'name': text,
                    'start_para': i
                })
        
        # 确定每个题型的结束位置
        for i in range(len(sections)):
            if i + 1 < len(sections):
                sections[i]['end_para'] = sections[i + 1]['start_para']
            else:
                sections[i]['end_para'] = len(doc.paragraphs)
        
        return sections
    
    def find_questions_in_section(self, doc: Document, section: Dict[str, Any]) -> List[Dict[str, Any]]:
        """提取章节内部的题目编号范围"""
        questions: List[Dict[str, Any]] = []
        current_start = section['start_para']
        last_boundary = section['start_para']
        dataset_heading_pattern = re.compile(r'^（[一二三四五六七八九十]+）$')
        in_preface = False
        # 分组ID：遇到（（一）（二）（三）…）时递增，用于资料分析等大题分组
        group_id = 0
        current_group_start = section['start_para']

        for i in range(section['start_para'], section['end_para']):
            if i >= len(doc.paragraphs):
                break
            text = doc.paragraphs[i].text.strip()

            if dataset_heading_pattern.match(text):
                if questions:
                    questions[-1]['end'] = last_boundary
                current_start = i
                last_boundary = i
                in_preface = True
                # 新的数据集开始，切换到新的组
                group_id += 1
                current_group_start = i
                continue

            if re.match(self.question_pattern, text):
                if questions:
                    # 结束上一题到当前数字标题前的最后一行
                    questions[-1]['end'] = last_boundary
                # 新题起点从上一段落边界开始，确保包含本题题干材料（位于题号前的图文）
                question_start = current_start if not questions else last_boundary
                questions.append({
                    'number': int(text),
                    'start': question_start,
                    'number_index': i,
                    'group_id': group_id,
                    'group_start': current_group_start
                })
                current_start = i
                last_boundary = i
                in_preface = False
            else:
                if not in_preface:
                    last_boundary = i + 1

        if questions:
            questions[-1]['end'] = max(last_boundary, questions[-1]['number_index'] + 1)

        return questions

    def extract_question_content(self, doc: Document, question: Dict[str, Any]) -> Dict[str, Any]:
        """提取题目的完整内容"""
        content = []
        total_images = 0
        total_text_length = 0
        # 选项识别：A/B/C/D + 标点（全角/半角）
        option_pattern = re.compile(r'^[A-DＡ-Ｄ][\s\t]*[、\.．\)]')
        
        for para_idx in range(question['start'], question['end']):
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                para_text = para.text.strip()
                
                # 统计图片（简化方法）
                para_images = 0
                for run in para.runs:
                    if hasattr(run._element, 'xml'):
                        xml_str = str(run._element.xml)
                        if '<w:drawing' in xml_str or '<pic:pic' in xml_str or '<v:imagedata' in xml_str:
                            para_images += 1
                
                # 选项段落的图片不计入题目材料图片总数
                if not option_pattern.match(para_text):
                    total_images += para_images
                total_text_length += len(para_text)
                
                if para_text or para_images > 0:
                    content.append({
                        'paragraph_index': para_idx,
                        'text': para_text,
                        'images_count': para_images,
                        'text_length': len(para_text)
                    })
        
        return {
            'paragraphs': content,
            'paragraph_count': len(content),
            'total_images': total_images,
            'total_text_length': total_text_length
        }
    
    def validate_extraction_results(self, questions: List[Dict], total_questions: int, total_images: int) -> Dict[str, Any]:
        """验证提取结果"""
        issues = []
        
        # 验证题目数量
        if total_questions != 135:
            issues.append(f"题目数量异常: 预期135道，实际{total_questions}道")
        
        # 验证题目编号连续性
        numbers = [q['number'] for q in questions]
        numbers.sort()
        for i in range(1, len(numbers)):
            if numbers[i] != numbers[i-1] + 1:
                issues.append(f"题目编号不连续: {numbers[i-1]} -> {numbers[i]}")
                break
        
        # 验证图片数量
        if total_images < 3000:
            issues.append(f"图片数量偏少: 预期3000+张，实际{total_images}张")
        
        return {
            'total_questions': total_questions,
            'total_images': total_images,
            'issues': issues,
            'success_rate': max(0, (135 - len(issues)) / 135 * 100)
        }


class QuestionImageExtractor:
    """题目图片提取器"""

    def __init__(self, output_dir: str = "images") -> None:
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def extract_images_from_question(self, doc: Document, question_data: Dict[str, Any], question_id: str) -> List[Dict[str, Any]]:
        """从题目内容中提取图片并保存到输出目录"""
        images: List[Dict[str, Any]] = []
        question_number = question_data.get('number', 0)

        paragraph_range = question_data.get('paragraph_range', '0-0')
        content_paras = question_data.get('content', {}).get('paragraphs', [])
        paragraph_indices = sorted({int(p.get('paragraph_index')) for p in content_paras if p.get('paragraph_index') is not None})

        if not paragraph_indices:
            try:
                start_para, end_para = map(int, paragraph_range.split('-'))
            except ValueError:
                start_para, end_para = 0, len(doc.paragraphs)
            paragraph_indices = list(range(start_para, min(end_para, len(doc.paragraphs))))

        image_index = 0
        # 选项识别：A/B/C/D + 标点（包含全角/半角）。仅当前面出现过题号后才生效
        option_pattern = re.compile(r'^[A-DＡ-Ｄ][\s\t]*[、\.．\)]')
        seen_question_number = False

        for para_idx in paragraph_indices:
            if not (0 <= para_idx < len(doc.paragraphs)):
                continue
            para = doc.paragraphs[para_idx]
            para_text = para.text.strip()
            # 在遇到题号所在段之后，再识别选项段落；避免把材料中带A/B/C标注误判为选项
            if not seen_question_number:
                number_idx = question_data.get('number_index')
                if number_idx is not None:
                    if para_idx == number_idx:
                        seen_question_number = True
                else:
                    # 若缺少题号索引则默认首段之后均视为选项
                    seen_question_number = True
            is_option_para = bool(option_pattern.match(para_text)) and seen_question_number

            # 先从 run 中提取
            seen_hashes: set = set()
            def _emit(img_bytes: bytes, ext: str):
                nonlocal image_index
                # 去重：按前16字节做简单哈希
                h = img_bytes[:16]
                if h in seen_hashes:
                    return
                seen_hashes.add(h)
                file_extension = ext if ext else '.png'
                if not file_extension.startswith('.'):
                    file_extension = f'.{file_extension}'
                filename = f"question_{question_number}_{image_index}_{uuid.uuid4().hex[:8]}{file_extension}"
                filepath = os.path.join(self.output_dir, filename)
                if self.save_image_data(img_bytes, filepath):
                    images.append({
                        'filename': filename,
                        'context_text': para_text[:100],
                        'paragraph_index': para_idx,
                        'position_in_question': image_index,
                        'image_type': 'option' if is_option_para else 'material'
                    })
                    image_index += 1

            for run in para.runs:
                image_payload = self.extract_image_from_run(run, doc)
                if image_payload:
                    _emit(image_payload[0], image_payload[1])

            # 兜底：直接扫描段落XML，提取 r:embed / r:id 并从文档关系取图
            try:
                import re as _re
                para_xml = str(getattr(para._element, 'xml', ''))
                rel_ids = _re.findall(r'(?:r:embed|r:id)="([^"]+)"', para_xml)
                for rel_id in rel_ids:
                    part = None
                    try:
                        part = doc.part.related_parts.get(rel_id)
                    except Exception:
                        part = None
                    if part is None:
                        continue
                    img_bytes = getattr(part, 'blob', None) or getattr(part, '_blob', None)
                    if not img_bytes:
                        continue
                    partname = str(getattr(part, 'partname', ''))
                    ext = os.path.splitext(partname)[1].lower()
                    if not ext:
                        content_type = getattr(part, 'content_type', '')
                        ext = self._extension_from_content_type(content_type)
                    _emit(img_bytes, ext)
            except Exception:
                pass

        return images

    def extract_image_from_run(self, run, doc: Document) -> Optional[Tuple[bytes, str]]:
        """从 run 元素获取图片的原始数据和扩展名"""
        try:
            element = run._element

            # 优先使用命名空间方式读取关系ID
            ns = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'v': 'urn:schemas-microsoft-com:vml',
                'o': 'urn:schemas-microsoft-com:office:office',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            }
            embed_ids = []
            try:
                embed_ids = element.xpath('.//a:blip/@r:embed', namespaces=ns)
            except Exception:
                embed_ids = []
            if not embed_ids:
                try:
                    embed_ids = element.xpath('.//v:imagedata/@r:id', namespaces=ns)
                except Exception:
                    embed_ids = []
            # 兜底：使用 local-name() 匹配，防止命名空间不一致
            if not embed_ids:
                embed_ids = element.xpath(".//*[local-name()='blip']/@*[local-name()='embed']")
            if not embed_ids:
                embed_ids = element.xpath(".//*[local-name()='imagedata']/@*[local-name()='id']")
            # 终极兜底：直接在XML字符串里用正则找 r:embed / r:id
            if not embed_ids:
                try:
                    import re as _re
                    xml_str = str(getattr(element, 'xml', ''))
                    m = _re.findall(r'(?:r:embed|r:id)="([^"]+)"', xml_str)
                    if m:
                        embed_ids = m
                except Exception:
                    pass
            if not embed_ids:
                return None

            rel_id = embed_ids[0]
            image_part = None
            # 优先从当前 run 所属 part 获取
            try:
                image_part = run.part.related_parts.get(rel_id)
            except Exception:
                image_part = None
            # 兜底：从整个文档 part 关系里找
            if image_part is None:
                try:
                    image_part = doc.part.related_parts.get(rel_id)
                except Exception:
                    image_part = None
            if image_part is None:
                return None

            # 兼容不同实现的属性名
            image_bytes = getattr(image_part, 'blob', None)
            if image_bytes is None:
                image_bytes = getattr(image_part, '_blob', None)
            if image_bytes is None:
                return None
            partname = str(getattr(image_part, 'partname', ''))
            extension = os.path.splitext(partname)[1].lower()
            if not extension:
                content_type = getattr(image_part, 'content_type', '')
                extension = self._extension_from_content_type(content_type)

            return image_bytes, extension
        except Exception as exc:
            logger.warning(f"提取图片失败: {exc}")
            return None

    @staticmethod
    def _extension_from_content_type(content_type: str) -> str:
        mapping = {
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/x-emf': '.emf',
            'image/x-wmf': '.wmf',
        }
        return mapping.get(content_type.lower(), '')

    def save_image_data(self, image_data: bytes, filepath: str) -> bool:
        """将图片字节写入文件"""
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            with open(filepath, 'wb') as f:
                f.write(image_data)
            return True
        except Exception as exc:
            logger.warning(f"保存图片失败: {exc}")
            return False

def print_extraction_summary(result: Dict[str, Any]) -> None:
    """打印提取结果摘要"""
    if not result.get('success'):
        print(f"❌ 提取失败: {result.get('error', '未知错误')}")
        return
    
    print("=" * 60)
    print("🎉 题目提取完成！")
    print("=" * 60)
    print(f"📊 总题目数: {result['total_questions']}道")
    print(f"🖼️ 总图片数: {result['total_images']}张")
    print(f"📚 题型数量: {len(result['sections'])}个")
    
    print("\n📋 题型分布:")
    for section_key, section_info in result['sections'].items():
        print(f"  {section_info['name'][:15]}...: {section_info['count']}道题, {section_info['total_images']}张图")
    
    validation = result.get('validation', {})
    issues = validation.get('issues', [])
    success_rate = validation.get('success_rate', 0)
    
    print(f"\n✅ 成功率: {success_rate:.1f}%")
    
    if issues:
        print(f"\n⚠️ 发现 {len(issues)} 个问题:")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("\n🎯 完美提取，无问题发现！")
    
    print("=" * 60)


