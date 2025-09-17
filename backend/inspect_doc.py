from docx import Document
import sys

path = r"..\题目\2025年国家公务员录用考试《行测》题（副省级网友回忆版）.docx"
start = int(sys.argv[1]) if len(sys.argv) > 1 else 799
end = int(sys.argv[2]) if len(sys.argv) > 2 else 820

doc = Document(path)
for i in range(start, end):
    text = doc.paragraphs[i].text.strip()
    images = 0
    for run in doc.paragraphs[i].runs:
        xml = getattr(run._element, 'xml', '')
        if isinstance(xml, str) and ('<w:drawing' in xml or '<pic:pic' in xml or '<v:imagedata' in xml):
            images += 1
    if text or images:
        print(i, text.encode('unicode_escape'), images)
